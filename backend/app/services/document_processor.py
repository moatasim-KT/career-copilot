"""
Enhanced document processing service with streaming support and robust text extraction.
"""

import io
import tempfile
import zipfile
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Dict, List, Optional, Union, Tuple, Any
import asyncio

import chardet
import pypdf
from docx import Document
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
import pdfplumber
import camelot

from ..core.exceptions import (
    DocumentProcessingError,
    InvalidFileTypeError,
    FileSizeError,
    SecurityError
)
from ..core.logging import get_logger
from ..core.file_handler import FileSecurityValidator
from ..monitoring.metrics_collector import get_metrics_collector
from ..utils.error_handler import get_error_handler, ErrorCategory, ErrorSeverity
from .contract_structure_analyzer import get_contract_structure_analyzer, ContractStructure

logger = get_logger(__name__)
metrics_collector = get_metrics_collector()


class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    UNKNOWN = "unknown"


class ProcessingStatus(Enum):
    """Document processing status."""
    SUCCESS = "success"
    PARTIAL = "partial"
    FAILED = "failed"
    UNSUPPORTED = "unsupported"


class OCRQuality(Enum):
    """OCR quality levels."""
    HIGH = "high"
    MEDIUM = "medium"
    LOW = "low"
    FAILED = "failed"


@dataclass
class TableData:
    """Extracted table data."""
    table_number: int
    page_number: int
    rows: List[List[str]]
    headers: Optional[List[str]]
    extraction_method: str
    confidence: float


@dataclass
class ProcessedDocument:
    """Result of document processing."""
    content: str
    document_type: DocumentType
    status: ProcessingStatus
    metadata: Dict[str, any]
    warnings: List[str]
    errors: List[str]
    processing_time_seconds: float
    contract_structure: Optional[ContractStructure] = None


class DocumentProcessingService:
    """Enhanced document processing service with streaming and security."""
    
    # Maximum content length to prevent memory issues
    MAX_CONTENT_LENGTH = 10 * 1024 * 1024  # 10MB
    
    # Chunk size for streaming processing
    CHUNK_SIZE = 64 * 1024  # 64KB
    
    def __init__(self):
        """Initialize document processor."""
        self.file_validator = FileSecurityValidator()
        self.structure_analyzer = get_contract_structure_analyzer()
        self._init_docling()
        self._check_tesseract()
    
    def _init_docling(self):
        """Initialize docling converter with MLX acceleration if available."""
        try:
            from docling.document_converter import DocumentConverter, PdfFormatOption
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.pipeline_options import PdfPipelineOptions
            
            # Check for MLX availability
            try:
                import mlx.core as mx
                logger.info("MLX acceleration available and enabled")
            except ImportError:
                logger.info("MLX not available, using CPU")
            
            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = True
            pipeline_options.do_table_structure = True
            
            self.docling_converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
                }
            )
            logger.info("Docling converter initialized")
        except Exception as e:
            logger.warning(f"Docling not available: {e}")
            self.docling_converter = None
    
    def _check_tesseract(self):
        """Check if tesseract is available."""
        try:
            pytesseract.get_tesseract_version()
            self.tesseract_available = True
            logger.info("Tesseract OCR available")
        except Exception:
            self.tesseract_available = False
            logger.info("Tesseract OCR not available")
    
    async def process_document(
        self, 
        file_content: Union[bytes, str], 
        filename: str,
        validate_security: bool = True,
        enable_ocr: bool = True,
        enable_table_extraction: bool = True
    ) -> ProcessedDocument:
        """
        Process document with comprehensive text extraction and security validation.
        
        Args:
            file_content: Document content as bytes or file path
            filename: Original filename
            validate_security: Whether to perform security validation
            
        Returns:
            ProcessedDocument with extracted content and metadata
        """
        import time
        start_time = time.time()
        
        logger.info(f"Starting document processing for: {filename}")
        
        # Initialize result
        result = ProcessedDocument(
            content="",
            document_type=DocumentType.UNKNOWN,
            status=ProcessingStatus.FAILED,
            metadata={"filename": filename},
            warnings=[],
            errors=[],
            processing_time_seconds=0.0
        )
        
        try:
            error_handler = get_error_handler()
            
            # Handle file path vs content
            if isinstance(file_content, str):
                # File path provided
                file_path = Path(file_content)
                if not file_path.exists():
                    raise DocumentProcessingError(
                        f"File not found: {file_content}",
                        processing_stage="file_access",
                        details={"file_path": file_content, "filename": filename}
                    )
                
                try:
                    with open(file_path, 'rb') as f:
                        content_bytes = f.read()
                except PermissionError as e:
                    raise DocumentProcessingError(
                        f"Permission denied accessing file: {filename}",
                        processing_stage="file_access",
                        details={"file_path": file_content, "filename": filename},
                        cause=e
                    )
                except Exception as e:
                    raise DocumentProcessingError(
                        f"Failed to read file: {filename}",
                        processing_stage="file_access",
                        details={"file_path": file_content, "filename": filename},
                        cause=e
                    )
            else:
                # Content bytes provided
                content_bytes = file_content
            
            # Validate file size
            if len(content_bytes) == 0:
                raise DocumentProcessingError(
                    "File is empty",
                    processing_stage="content_validation",
                    details={"filename": filename, "file_size": 0}
                )
            
            if len(content_bytes) > self.MAX_CONTENT_LENGTH:
                raise FileSizeError(
                    f"File size exceeds maximum limit",
                    file_size=len(content_bytes),
                    max_size=self.MAX_CONTENT_LENGTH,
                    details={"filename": filename}
                )
            
            # Security validation
            if validate_security:
                try:
                    validation_result = self.file_validator.validate_file(content_bytes, filename)
                    logger.info(f"File validation successful for {filename}")
                except Exception as e:
                    raise SecurityError(
                        f"File failed security validation: {str(e)}",
                        violation_type="file_validation",
                        details={"filename": filename, "validation_error": str(e)},
                        cause=e
                    )
            
            # Content validation
            content_valid, content_warnings = await self.validate_document_content(content_bytes, filename)
            if not content_valid:
                result.errors.extend(content_warnings)
                return result
            else:
                result.warnings.extend(content_warnings)
            
            # Determine document type
            try:
                doc_type = self._determine_document_type(content_bytes, filename)
                result.document_type = doc_type
            except Exception as e:
                raise DocumentProcessingError(
                    f"Failed to determine document type for '{filename}'",
                    processing_stage="type_detection",
                    details={"filename": filename},
                    cause=e
                )
            
            # Process based on document type
            try:
                if doc_type == DocumentType.PDF:
                    await self._process_pdf(content_bytes, result)
                elif doc_type == DocumentType.DOCX:
                    await self._process_docx(content_bytes, result)
                elif doc_type == DocumentType.TXT:
                    await self._process_txt(content_bytes, result)
                else:
                    raise InvalidFileTypeError(
                        f"Unsupported document type: {doc_type.value}",
                        file_type=doc_type.value,
                        supported_types=["pdf", "docx", "txt"],
                        details={"filename": filename}
                    )
            except (InvalidFileTypeError, DocumentProcessingError):
                # Re-raise known exceptions
                raise
            except Exception as e:
                raise DocumentProcessingError(
                    f"Failed to process {doc_type.value} document '{filename}'",
                    processing_stage="content_extraction",
                    details={"filename": filename, "document_type": doc_type.value},
                    cause=e
                )
            
            # Validate content length
            if len(result.content) > self.MAX_CONTENT_LENGTH:
                result.warnings.append(f"Content truncated to {self.MAX_CONTENT_LENGTH} characters")
                result.content = result.content[:self.MAX_CONTENT_LENGTH]
                result.status = ProcessingStatus.PARTIAL
            
            # Try OCR if minimal text extracted
            if result.document_type == DocumentType.PDF and enable_ocr:
                if len(result.content.strip()) < 100:
                    logger.info(f"Minimal text, attempting OCR for {filename}")
                    await self._process_with_ocr(content_bytes, result)
            
            # Try docling for enhanced extraction
            if self.docling_converter and result.document_type == DocumentType.PDF:
                await self._process_with_docling(content_bytes, result)
            
            # Extract tables
            if enable_table_extraction and result.document_type == DocumentType.PDF:
                tables = await self._extract_tables(content_bytes, result)
                if tables:
                    result.metadata['tables'] = [
                        {
                            'table_number': t.table_number,
                            'page_number': t.page_number,
                            'row_count': len(t.rows),
                            'column_count': len(t.rows[0]) if t.rows else 0,
                            'extraction_method': t.extraction_method,
                            'confidence': t.confidence
                        }
                        for t in tables
                    ]
                    result.metadata['table_count'] = len(tables)
            
            # Analyze contract structure if content is available
            if result.content.strip():
                try:
                    result.contract_structure = self.structure_analyzer.analyze_structure(result.content)
                    logger.info(f"Contract structure analysis completed: {len(result.contract_structure.sections)} sections, {len(result.contract_structure.clauses)} clauses")
                except Exception as e:
                    logger.warning(f"Contract structure analysis failed: {e}")
                    result.warnings.append(f"Structure analysis failed: {str(e)}")
            
            # Final processing
            result.processing_time_seconds = time.time() - start_time
            
            if result.content.strip():
                if result.status != ProcessingStatus.PARTIAL:
                    result.status = ProcessingStatus.SUCCESS
                logger.info(f"Document processing completed: {filename} ({len(result.content)} chars)")
            else:
                result.errors.append("No text content extracted from document")
                result.status = ProcessingStatus.FAILED
            
            # Record file processing metrics
            metrics_collector.record_file_upload(
                file_type=doc_type.value,
                status=result.status.value,
                file_size=len(content_bytes),
                processing_duration=result.processing_time_seconds
            )
            
            return result
            
        except (DocumentProcessingError, InvalidFileTypeError, FileSizeError, SecurityError) as e:
            # Re-raise known application errors
            result.processing_time_seconds = time.time() - start_time
            
            # Record failed file processing metrics
            metrics_collector.record_file_upload(
                file_type=result.document_type.value,
                status="failed",
                file_size=len(content_bytes) if 'content_bytes' in locals() else 0,
                processing_duration=result.processing_time_seconds
            )
            
            raise
            
        except Exception as e:
            # Handle unexpected errors
            error_handler = get_error_handler()
            result.processing_time_seconds = time.time() - start_time
            
            error_context = error_handler.handle_error(
                error=e,
                category=ErrorCategory.FILE_PROCESSING,
                severity=ErrorSeverity.HIGH,
                user_message=f"Unexpected error processing document '{filename}'",
                technical_details={
                    "filename": filename,
                    "processing_time": result.processing_time_seconds,
                    "error_type": type(e).__name__
                }
            )
            
            # Record failed file processing metrics
            metrics_collector.record_file_upload(
                file_type=result.document_type.value,
                status="failed",
                file_size=len(content_bytes) if 'content_bytes' in locals() else 0,
                processing_duration=result.processing_time_seconds
            )
            
            raise DocumentProcessingError(
                f"Unexpected error processing document '{filename}': {str(e)}",
                processing_stage="unknown",
                details={"filename": filename, "error_id": error_context.error_id},
                cause=e
            )
            file_size = len(content_bytes) if 'content_bytes' in locals() else 0
            doc_type = result.document_type.value if result.document_type != DocumentType.UNKNOWN else "unknown"
            
            metrics_collector.record_file_upload(
                file_type=doc_type,
                status="failed",
                file_size=file_size,
                processing_duration=result.processing_time_seconds
            )
            
            return result
    
    def _determine_document_type(self, content: bytes, filename: str) -> DocumentType:
        """Determine document type from content and filename with enhanced detection."""
        if not content:
            raise DocumentProcessingError(
                "Cannot determine document type of empty file",
                processing_stage="type_detection",
                details={"filename": filename, "content_length": 0}
            )
        
        try:
            # Use the enhanced MIME type detection from file validator
            detected_mime = self.file_validator._detect_mime_type_robust(content, filename)
            
            # Map MIME types to document types
            mime_to_doc_type = {
                'application/pdf': DocumentType.PDF,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentType.DOCX,
                'text/plain': DocumentType.TXT,
                'text/csv': DocumentType.TXT,
                'application/msword': DocumentType.UNKNOWN,  # Legacy .doc not supported
                'application/rtf': DocumentType.UNKNOWN,     # RTF not supported
                'text/html': DocumentType.UNKNOWN,          # HTML not supported for job application tracking
                'application/xml': DocumentType.UNKNOWN,    # XML not supported
                'application/zip': DocumentType.UNKNOWN,    # Generic ZIP not supported
                'application/octet-stream': DocumentType.UNKNOWN
            }
            
            doc_type = mime_to_doc_type.get(detected_mime, DocumentType.UNKNOWN)
            
            if doc_type == DocumentType.UNKNOWN:
                # Provide specific error messages for unsupported but recognized formats
                if detected_mime == 'application/msword':
                    raise InvalidFileTypeError(
                        "Legacy Word documents (.doc) are not supported. Please convert to .docx format.",
                        file_type="doc",
                        supported_types=["pdf", "docx", "txt"],
                        details={"filename": filename, "detected_mime": detected_mime}
                    )
                elif detected_mime == 'text/html':
                    raise InvalidFileTypeError(
                        "HTML files are not supported for job application tracking. Please use PDF, DOCX, or plain text.",
                        file_type="html",
                        supported_types=["pdf", "docx", "txt"],
                        details={"filename": filename, "detected_mime": detected_mime}
                    )
                elif detected_mime == 'application/rtf':
                    raise InvalidFileTypeError(
                        "RTF files are not supported. Please convert to PDF, DOCX, or plain text.",
                        file_type="rtf",
                        supported_types=["pdf", "docx", "txt"],
                        details={"filename": filename, "detected_mime": detected_mime}
                    )
                else:
                    raise InvalidFileTypeError(
                        f"Unsupported file type detected: {detected_mime}",
                        file_type=detected_mime,
                        supported_types=["pdf", "docx", "txt"],
                        details={"filename": filename, "detected_mime": detected_mime}
                    )
            
            logger.info(f"Document type determined: {filename} -> {doc_type.value} (MIME: {detected_mime})")
            return doc_type
            
        except (InvalidFileTypeError, DocumentProcessingError):
            # Re-raise known exceptions
            raise
        except Exception as e:
            # Handle unexpected errors in type detection
            raise DocumentProcessingError(
                f"Failed to determine document type for '{filename}': {str(e)}",
                processing_stage="type_detection",
                details={"filename": filename, "error_type": type(e).__name__},
                cause=e
            )
    
    async def _process_pdf(self, content: bytes, result: ProcessedDocument) -> None:
        """Process PDF document with streaming support and enhanced error handling."""
        try:
            # Create PDF reader from bytes
            pdf_stream = io.BytesIO(content)
            
            try:
                pdf_reader = pypdf.PdfReader(pdf_stream)
            except Exception as e:
                result.errors.append(f"Failed to open PDF: {str(e)}")
                return
            
            # Extract metadata
            metadata = {
                "page_count": len(pdf_reader.pages),
                "encrypted": pdf_reader.is_encrypted,
                "pdf_version": getattr(pdf_reader, 'pdf_version', 'unknown')
            }
            
            # Handle encrypted PDFs
            if pdf_reader.is_encrypted:
                result.warnings.append("PDF is encrypted - attempting to decrypt")
                try:
                    # Try empty password first
                    if not pdf_reader.decrypt(''):
                        result.errors.append("PDF is password protected and cannot be processed")
                        return
                    else:
                        result.warnings.append("PDF decrypted successfully with empty password")
                except Exception as e:
                    result.errors.append(f"Failed to decrypt PDF: {str(e)}")
                    return
            
            # Check if PDF has pages
            if len(pdf_reader.pages) == 0:
                result.errors.append("PDF contains no pages")
                return
            
            # Extract text from all pages with streaming
            text_parts = []
            total_chars = 0
            pages_processed = 0
            pages_with_text = 0
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    pages_processed += 1
                    
                    if page_text and page_text.strip():
                        pages_with_text += 1
                        
                        # Clean up the extracted text
                        page_text = page_text.strip()
                        
                        # Check content length limit
                        if total_chars + len(page_text) > self.MAX_CONTENT_LENGTH:
                            remaining = self.MAX_CONTENT_LENGTH - total_chars
                            if remaining > 0:
                                text_parts.append(page_text[:remaining])
                            result.warnings.append(f"Content truncated at page {page_num + 1} due to size limit")
                            result.status = ProcessingStatus.PARTIAL
                            break
                        
                        text_parts.append(page_text)
                        total_chars += len(page_text)
                    
                except Exception as e:
                    result.warnings.append(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            # Update metadata with processing stats
            metadata.update({
                "pages_processed": pages_processed,
                "pages_with_text": pages_with_text,
                "text_extraction_rate": pages_with_text / pages_processed if pages_processed > 0 else 0
            })
            
            # Combine text
            result.content = '\n\n'.join(text_parts)  # Use double newlines for page breaks
            
            # Check if any text was extracted
            if not result.content.strip():
                result.warnings.append("No text content could be extracted from PDF - document may be image-based")
                result.errors.append("PDF appears to contain no extractable text")
                return
            
            # Add PDF-specific metadata
            if hasattr(pdf_reader, 'metadata') and pdf_reader.metadata:
                try:
                    pdf_metadata = pdf_reader.metadata
                    metadata.update({
                        "title": str(pdf_metadata.get('/Title', '')),
                        "author": str(pdf_metadata.get('/Author', '')),
                        "subject": str(pdf_metadata.get('/Subject', '')),
                        "creator": str(pdf_metadata.get('/Creator', '')),
                        "producer": str(pdf_metadata.get('/Producer', '')),
                        "creation_date": str(pdf_metadata.get('/CreationDate', '')),
                        "modification_date": str(pdf_metadata.get('/ModDate', ''))
                    })
                except Exception as e:
                    result.warnings.append(f"Failed to extract PDF metadata: {str(e)}")
            
            result.metadata.update(metadata)
            logger.info(f"PDF processing completed: {pages_with_text}/{pages_processed} pages with text, {len(result.content)} characters extracted")
            
        except Exception as e:
            result.errors.append(f"PDF processing error: {str(e)}")
            logger.error(f"PDF processing failed: {e}", exc_info=True)
            raise
    
    async def _process_docx(self, content: bytes, result: ProcessedDocument) -> None:
        """Process DOCX document with streaming support and enhanced error handling."""
        temp_file_path = None
        try:
            # Validate DOCX file structure first
            if not content.startswith(b'PK'):
                result.errors.append("Invalid DOCX file format - not a valid ZIP archive")
                return
            
            # Create temporary file for python-docx
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                # Open document
                doc = Document(temp_file_path)
                
                # Extract metadata
                metadata = {
                    "paragraph_count": len(doc.paragraphs),
                    "has_tables": len(doc.tables) > 0,
                    "table_count": len(doc.tables)
                }
                
                # Extract core properties if available
                if hasattr(doc, 'core_properties'):
                    try:
                        core_props = doc.core_properties
                        metadata.update({
                            "title": str(core_props.title or ''),
                            "author": str(core_props.author or ''),
                            "subject": str(core_props.subject or ''),
                            "created": str(core_props.created) if core_props.created else '',
                            "modified": str(core_props.modified) if core_props.modified else '',
                            "last_modified_by": str(core_props.last_modified_by or '')
                        })
                    except Exception as e:
                        result.warnings.append(f"Failed to extract document properties: {str(e)}")
                
                # Extract text with streaming approach
                text_parts = []
                total_chars = 0
                paragraphs_processed = 0
                paragraphs_with_text = 0
                
                # Extract paragraph text
                for para_num, paragraph in enumerate(doc.paragraphs):
                    try:
                        para_text = paragraph.text
                        paragraphs_processed += 1
                        
                        if para_text and para_text.strip():
                            paragraphs_with_text += 1
                            para_text = para_text.strip()
                            
                            # Check content length limit
                            if total_chars + len(para_text) > self.MAX_CONTENT_LENGTH:
                                remaining = self.MAX_CONTENT_LENGTH - total_chars
                                if remaining > 0:
                                    text_parts.append(para_text[:remaining])
                                result.warnings.append(f"Content truncated at paragraph {para_num + 1} due to size limit")
                                result.status = ProcessingStatus.PARTIAL
                                break
                            
                            text_parts.append(para_text)
                            total_chars += len(para_text)
                    
                    except Exception as e:
                        result.warnings.append(f"Failed to extract text from paragraph {para_num + 1}: {str(e)}")
                        continue
                
                # Extract table text if not at limit
                tables_processed = 0
                tables_with_text = 0
                
                if total_chars < self.MAX_CONTENT_LENGTH and doc.tables:
                    for table_num, table in enumerate(doc.tables):
                        try:
                            tables_processed += 1
                            table_text = self._extract_table_text(table)
                            
                            if table_text and table_text.strip():
                                tables_with_text += 1
                                
                                # Check content length limit
                                if total_chars + len(table_text) > self.MAX_CONTENT_LENGTH:
                                    remaining = self.MAX_CONTENT_LENGTH - total_chars
                                    if remaining > 0:
                                        text_parts.append(f"\n[Table {table_num + 1}]\n{table_text[:remaining]}")
                                    result.warnings.append(f"Content truncated at table {table_num + 1} due to size limit")
                                    result.status = ProcessingStatus.PARTIAL
                                    break
                                
                                text_parts.append(f"\n[Table {table_num + 1}]\n{table_text}")
                                total_chars += len(table_text)
                        
                        except Exception as e:
                            result.warnings.append(f"Failed to extract text from table {table_num + 1}: {str(e)}")
                            continue
                
                # Update metadata with processing stats
                metadata.update({
                    "paragraphs_processed": paragraphs_processed,
                    "paragraphs_with_text": paragraphs_with_text,
                    "tables_processed": tables_processed,
                    "tables_with_text": tables_with_text,
                    "text_extraction_rate": paragraphs_with_text / paragraphs_processed if paragraphs_processed > 0 else 0
                })
                
                # Combine text
                result.content = '\n'.join(text_parts)
                
                # Check if any text was extracted
                if not result.content.strip():
                    result.warnings.append("No text content could be extracted from DOCX document")
                    result.errors.append("DOCX document appears to contain no extractable text")
                    return
                
                result.metadata.update(metadata)
                logger.info(f"DOCX processing completed: {paragraphs_with_text}/{paragraphs_processed} paragraphs, {tables_with_text}/{tables_processed} tables, {len(result.content)} characters extracted")
                
            except Exception as e:
                result.errors.append(f"Failed to process DOCX document: {str(e)}")
                logger.error(f"DOCX processing failed: {e}", exc_info=True)
                return
                
        except Exception as e:
            result.errors.append(f"DOCX processing error: {str(e)}")
            logger.error(f"DOCX processing failed: {e}", exc_info=True)
            raise
            
        finally:
            # Clean up temporary file
            if temp_file_path and Path(temp_file_path).exists():
                try:
                    Path(temp_file_path).unlink()
                except Exception as e:
                    logger.warning(f"Failed to clean up temporary file {temp_file_path}: {e}")
    
    def _extract_table_text(self, table) -> str:
        """Extract text from DOCX table."""
        table_text = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_text.append(cell_text)
            
            if any(cell.strip() for cell in row_text):  # Skip empty rows
                table_text.append('\t'.join(row_text))
        
        return '\n'.join(table_text)
    
    async def _process_txt(self, content: bytes, result: ProcessedDocument) -> None:
        """Process text document with enhanced encoding detection and validation."""
        try:
            # Check if content is empty
            if not content:
                result.errors.append("Text file is empty")
                return
            
            # Detect encoding
            encoding_result = chardet.detect(content)
            detected_encoding = encoding_result.get('encoding', 'utf-8')
            confidence = encoding_result.get('confidence', 0.0)
            
            # Try detected encoding first
            text_content = None
            encoding_used = None
            decoding_errors = []
            
            if detected_encoding and confidence > 0.7:
                try:
                    text_content = content.decode(detected_encoding)
                    encoding_used = detected_encoding
                    logger.debug(f"Successfully decoded text using detected encoding: {detected_encoding} (confidence: {confidence:.2f})")
                except UnicodeDecodeError as e:
                    decoding_errors.append(f"Failed to decode with {detected_encoding}: {str(e)}")
            
            # Fallback encodings
            if text_content is None:
                fallback_encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'ascii']
                
                for encoding in fallback_encodings:
                    try:
                        text_content = content.decode(encoding)
                        encoding_used = encoding
                        if encoding != detected_encoding:
                            result.warnings.append(f"Used fallback encoding: {encoding} (detected: {detected_encoding})")
                        logger.debug(f"Successfully decoded text using fallback encoding: {encoding}")
                        break
                    except UnicodeDecodeError as e:
                        decoding_errors.append(f"Failed to decode with {encoding}: {str(e)}")
                        continue
            
            if text_content is None:
                # Last resort - decode with errors='replace'
                text_content = content.decode('utf-8', errors='replace')
                encoding_used = 'utf-8 (with character replacement)'
                result.warnings.append("Text decoded with character replacement due to encoding issues")
                result.warnings.extend(decoding_errors[:3])  # Include first few decoding errors
            
            # Validate and clean text content
            if not text_content.strip():
                result.errors.append("Text file contains no readable content")
                return
            
            # Remove null bytes and other problematic characters
            text_content = text_content.replace('\x00', '')
            
            # Check for binary content indicators
            binary_indicators = ['\x01', '\x02', '\x03', '\x04', '\x05', '\x06', '\x07', '\x08']
            if any(indicator in text_content for indicator in binary_indicators):
                result.warnings.append("File may contain binary data - some content may not display correctly")
            
            # Check content length
            original_length = len(text_content)
            if original_length > self.MAX_CONTENT_LENGTH:
                result.warnings.append(f"Content truncated from {original_length} to {self.MAX_CONTENT_LENGTH} characters")
                text_content = text_content[:self.MAX_CONTENT_LENGTH]
                result.status = ProcessingStatus.PARTIAL
            
            result.content = text_content
            
            # Calculate text statistics
            lines = text_content.split('\n')
            words = text_content.split()
            
            # Estimate content quality
            avg_line_length = sum(len(line) for line in lines) / len(lines) if lines else 0
            non_empty_lines = [line for line in lines if line.strip()]
            
            # Add text-specific metadata
            metadata = {
                "encoding": encoding_used,
                "encoding_confidence": confidence,
                "detected_encoding": detected_encoding,
                "line_count": len(lines),
                "non_empty_line_count": len(non_empty_lines),
                "character_count": len(text_content),
                "word_count": len(words),
                "average_line_length": round(avg_line_length, 2),
                "content_density": len(non_empty_lines) / len(lines) if lines else 0,
                "original_size_bytes": len(content),
                "processed_size_chars": len(text_content)
            }
            
            result.metadata.update(metadata)
            logger.info(f"Text processing completed: {len(words)} words, {len(lines)} lines, encoding: {encoding_used}")
            
        except Exception as e:
            result.errors.append(f"Text processing error: {str(e)}")
            logger.error(f"Text processing failed: {e}", exc_info=True)
            raise
    
    async def extract_metadata_only(self, file_content: bytes, filename: str) -> Dict[str, any]:
        """Extract only metadata without full text processing."""
        doc_type = self._determine_document_type(file_content, filename)
        metadata = {"filename": filename, "document_type": doc_type.value}
        
        try:
            if doc_type == DocumentType.PDF:
                pdf_stream = io.BytesIO(file_content)
                pdf_reader = pypdf.PdfReader(pdf_stream)
                metadata.update({
                    "page_count": len(pdf_reader.pages),
                    "encrypted": pdf_reader.is_encrypted,
                    "file_size": len(file_content)
                })
                
            elif doc_type == DocumentType.DOCX:
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name
                
                try:
                    doc = Document(temp_file_path)
                    metadata.update({
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables),
                        "file_size": len(file_content)
                    })
                finally:
                    Path(temp_file_path).unlink(missing_ok=True)
                    
            elif doc_type == DocumentType.TXT:
                encoding_result = chardet.detect(file_content)
                metadata.update({
                    "encoding": encoding_result.get('encoding', 'unknown'),
                    "encoding_confidence": encoding_result.get('confidence', 0.0),
                    "file_size": len(file_content)
                })
            
        except Exception as e:
            logger.warning(f"Failed to extract metadata from {filename}: {e}")
            metadata["extraction_error"] = str(e)
        
        return metadata
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats."""
        return [doc_type.value for doc_type in DocumentType if doc_type != DocumentType.UNKNOWN]
    
    def validate_document_size(self, file_size: int, doc_type: DocumentType) -> Tuple[bool, Optional[str]]:
        """Validate document size against limits with detailed feedback."""
        size_limits = {
            DocumentType.PDF: 50 * 1024 * 1024,    # 50MB
            DocumentType.DOCX: 25 * 1024 * 1024,   # 25MB
            DocumentType.TXT: 10 * 1024 * 1024,    # 10MB
        }
        
        limit = size_limits.get(doc_type, 10 * 1024 * 1024)
        
        if file_size <= 0:
            return False, "File is empty or has invalid size"
        
        if file_size > limit:
            limit_mb = limit / (1024 * 1024)
            file_mb = file_size / (1024 * 1024)
            return False, f"File size ({file_mb:.1f}MB) exceeds limit for {doc_type.value} files ({limit_mb:.0f}MB)"
        
        # Warn about very small files
        if file_size < 100:  # Less than 100 bytes
            return True, f"Warning: File is very small ({file_size} bytes) and may not contain meaningful content"
        
        return True, None
    
    async def validate_document_content(self, content: bytes, filename: str) -> Tuple[bool, List[str]]:
        """Validate document content for common issues."""
        warnings = []
        
        if not content:
            return False, ["Document is empty"]
        
        # Check file size
        doc_type = self._determine_document_type(content, filename)
        is_valid, size_message = self.validate_document_size(len(content), doc_type)
        
        if not is_valid:
            return False, [size_message]
        elif size_message:
            warnings.append(size_message)
        
        # Check for unsupported document types
        if doc_type == DocumentType.UNKNOWN:
            return False, ["Unsupported document type - only PDF, DOCX, and TXT files are supported"]
        
        # Type-specific validation
        if doc_type == DocumentType.PDF:
            if not content.startswith(b'%PDF-'):
                return False, ["File appears to be corrupted - invalid PDF header"]
            
            # Check for PDF version
            try:
                header = content[:20].decode('ascii', errors='ignore')
                if 'PDF-1.' in header:
                    version = header.split('PDF-1.')[1][:1]
                    if version.isdigit() and int(version) < 3:
                        warnings.append(f"Old PDF version detected (1.{version}) - may have compatibility issues")
            except Exception:
                pass
        
        elif doc_type == DocumentType.DOCX:
            if not content.startswith(b'PK'):
                return False, ["File appears to be corrupted - invalid DOCX format"]
            
            # Check if it's actually a DOCX and not just any ZIP file
            if b'word/' not in content[:2048] and not filename.lower().endswith('.docx'):
                return False, ["File is not a valid DOCX document"]
        
        elif doc_type == DocumentType.TXT:
            # Check for binary content in text files
            try:
                # Sample first 1KB
                sample = content[:1024]
                sample.decode('utf-8')
            except UnicodeDecodeError:
                # Try other encodings
                decoded = False
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        sample.decode(encoding)
                        decoded = True
                        warnings.append(f"Text file uses {encoding} encoding instead of UTF-8")
                        break
                    except UnicodeDecodeError:
                        continue
                
                if not decoded:
                    return False, ["Text file contains binary data or uses unsupported encoding"]
        
        return True, warnings
    
    async def _detect_scanned_pdf(self, content: bytes) -> Tuple[bool, float]:
        """Detect if PDF is scanned."""
        try:
            pdf_stream = io.BytesIO(content)
            pdf_reader = pypdf.PdfReader(pdf_stream)
            
            if len(pdf_reader.pages) == 0:
                return False, 0.0
            
            sample_pages = min(3, len(pdf_reader.pages))
            text_chars = 0
            
            for i in range(sample_pages):
                page_text = pdf_reader.pages[i].extract_text()
                text_chars += len(page_text.strip())
            
            avg_chars_per_page = text_chars / sample_pages
            
            if avg_chars_per_page < 50:
                return True, 0.9
            elif avg_chars_per_page < 200:
                return True, 0.6
            else:
                return False, 0.3
        except Exception as e:
            logger.warning(f"Failed to detect scanned PDF: {e}")
            return False, 0.0
    
    async def _process_with_ocr(self, content: bytes, result: ProcessedDocument) -> None:
        """Process PDF with OCR."""
        if not self.tesseract_available:
            result.warnings.append("OCR requested but Tesseract not available")
            return
        
        try:
            is_scanned, confidence = await self._detect_scanned_pdf(content)
            result.metadata['is_scanned'] = is_scanned
            result.metadata['scanned_confidence'] = confidence
            
            if not is_scanned and confidence < 0.5:
                return
            
            logger.info(f"Processing scanned PDF with OCR (confidence: {confidence})")
            
            images = await asyncio.to_thread(convert_from_bytes, content, dpi=300, fmt='png')
            
            ocr_texts = []
            ocr_quality_scores = []
            
            for page_num, image in enumerate(images, 1):
                try:
                    page_text = await asyncio.to_thread(
                        pytesseract.image_to_string, image, lang='eng', config='--psm 1'
                    )
                    
                    ocr_data = await asyncio.to_thread(
                        pytesseract.image_to_data, image, output_type=pytesseract.Output.DICT
                    )
                    
                    confidences = [int(conf) for conf in ocr_data['conf'] if conf != '-1']
                    avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                    
                    ocr_texts.append(page_text)
                    ocr_quality_scores.append(avg_confidence)
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num}: {e}")
                    ocr_texts.append("")
                    ocr_quality_scores.append(0)
            
            ocr_text = '\n\n'.join(ocr_texts)
            avg_quality = sum(ocr_quality_scores) / len(ocr_quality_scores) if ocr_quality_scores else 0
            
            if avg_quality >= 80:
                ocr_quality = OCRQuality.HIGH
            elif avg_quality >= 60:
                ocr_quality = OCRQuality.MEDIUM
            elif avg_quality >= 40:
                ocr_quality = OCRQuality.LOW
            else:
                ocr_quality = OCRQuality.FAILED
            
            if len(ocr_text.strip()) > len(result.content.strip()):
                result.content = ocr_text
                result.metadata['extraction_method'] = 'ocr'
                result.metadata['ocr_quality'] = ocr_quality.value
                result.metadata['ocr_confidence'] = avg_quality
                result.metadata['ocr_pages_processed'] = len(images)
                
                if ocr_quality in [OCRQuality.LOW, OCRQuality.FAILED]:
                    result.warnings.append(f"OCR quality is {ocr_quality.value}")
                
                logger.info(f"OCR completed: {len(ocr_text)} chars, quality: {ocr_quality.value}")
        except Exception as e:
            logger.error(f"OCR processing failed: {e}")
            result.warnings.append(f"OCR processing failed: {str(e)}")
    
    async def _process_with_docling(self, content: bytes, result: ProcessedDocument) -> None:
        """Process with docling for enhanced extraction."""
        try:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            
            try:
                doc_result = await asyncio.to_thread(self.docling_converter.convert, tmp_path)
                docling_text = doc_result.document.export_to_markdown()
                
                if len(docling_text.strip()) > len(result.content.strip()):
                    result.content = docling_text
                    result.metadata['extraction_method'] = 'docling'
                    logger.info(f"Docling extraction: {len(docling_text)} chars")
                
                if hasattr(doc_result.document, 'tables'):
                    result.metadata['docling_tables'] = len(doc_result.document.tables)
            finally:
                Path(tmp_path).unlink(missing_ok=True)
        except Exception as e:
            logger.warning(f"Docling processing failed: {e}")
    
    async def _extract_tables(self, content: bytes, result: ProcessedDocument) -> List[TableData]:
        """Extract tables from PDF."""
        tables = []
        
        try:
            tables.extend(await self._extract_tables_pdfplumber(content))
        except Exception as e:
            logger.warning(f"pdfplumber table extraction failed: {e}")
        
        if len(tables) == 0:
            try:
                tables.extend(await self._extract_tables_camelot(content))
            except Exception as e:
                logger.warning(f"camelot table extraction failed: {e}")
        
        return tables
    
    async def _extract_tables_pdfplumber(self, content: bytes) -> List[TableData]:
        """Extract tables using pdfplumber."""
        tables = []
        
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_tables = page.extract_tables()
                    
                    for table_num, table in enumerate(page_tables, 1):
                        if table and len(table) > 0:
                            headers = table[0] if len(table) > 1 else None
                            rows = table[1:] if len(table) > 1 else table
                            
                            tables.append(TableData(
                                table_number=len(tables) + 1,
                                page_number=page_num,
                                rows=rows,
                                headers=headers,
                                extraction_method='pdfplumber',
                                confidence=0.8
                            ))
            
            logger.info(f"Extracted {len(tables)} tables with pdfplumber")
        except Exception as e:
            logger.error(f"pdfplumber extraction failed: {e}")
            raise
        
        return tables
    
    async def _extract_tables_camelot(self, content: bytes) -> List[TableData]:
        """Extract tables using camelot."""
        tables = []
        
        try:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            
            try:
                camelot_tables = await asyncio.to_thread(
                    camelot.read_pdf, tmp_path, pages='all', flavor='lattice'
                )
                
                for table_num, table in enumerate(camelot_tables, 1):
                    df = table.df
                    rows = df.values.tolist()
                    headers = df.columns.tolist() if not df.columns.equals(range(len(df.columns))) else None
                    
                    tables.append(TableData(
                        table_number=table_num,
                        page_number=table.page,
                        rows=rows,
                        headers=headers,
                        extraction_method='camelot',
                        confidence=table.accuracy / 100.0
                    ))
                
                logger.info(f"Extracted {len(tables)} tables with camelot")
            finally:
                Path(tmp_path).unlink(missing_ok=True)
        except Exception as e:
            logger.error(f"camelot extraction failed: {e}")
            raise
        
        return tables


# Global service instance
_document_processor_service: Optional[DocumentProcessingService] = None


def get_document_processor_service() -> DocumentProcessingService:
    """Get or create the global document processor service instance."""
    global _document_processor_service
    
    if _document_processor_service is None:
        _document_processor_service = DocumentProcessingService()
    
    return _document_processor_service


# Compatibility alias
get_document_processor = get_document_processor_service